"""
DOCX generator for images with filenames displayed below each image
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from pathlib import Path


def set_cell_border(cell, **kwargs):
    """
    Set cell border for a table cell
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_val = kwargs.get(border_name)
        if border_val:
            tag = 'w:{}'.format(border_name)
            element = OxmlElement(tag)
            for key, value in border_val.items():
                element.set(qn('w:{}'.format(key)), str(value))
            tcPr.append(element)


def create_image_docx(images, image_names, output_path, images_per_page=2, quality=95, max_width=None, max_height=None):
    """
    Create DOCX document with images and filenames displayed below each image
    
    Args:
        images: List of image file paths (in order)
        image_names: List of filenames (without extension) to display below each image
        output_path: Path to save DOCX file
        images_per_page: Number of images per page (determines grid layout)
        quality: JPEG quality (1-100) - used when processing images
        max_width: Max image width in pixels (None = original)
        max_height: Max image height in pixels (None = original)
    """
    import tempfile
    import os
    
    doc = Document()
    
    # Set default margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Determine grid layout
    if images_per_page == 1:
        cols = 1
    elif images_per_page == 2:
        cols = 2
    elif images_per_page == 4:
        cols = 2  # 2x2
    elif images_per_page == 6:
        cols = 3  # 3x2
    elif images_per_page == 9:
        cols = 3  # 3x3
    elif images_per_page == 12:
        cols = 3  # 3x4
    else:
        cols = 2  # Default
    
    # Calculate column and image width
    page_width = Inches(8.27)
    margins = Inches(0.5) * 2  # Left + right
    usable_width = page_width - margins
    column_width = usable_width / cols
    image_width = Inches(3.3) if cols == 2 else Inches(2.2) if cols == 3 else Inches(7.0)
    
    # Process images if resizing needed
    processed_images = []
    temp_files = []
    
    for img_path in images:
        try:
            img_path_obj = Path(img_path)
            if not img_path_obj.exists():
                print(f"Warning: Image not found: {img_path}")
                continue
            
            pil_img = Image.open(img_path)
            
            # Resize if needed
            if max_width or max_height:
                if max_width and max_height:
                    pil_img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                elif max_width:
                    ratio = max_width / pil_img.width
                    new_height = int(pil_img.height * ratio)
                    pil_img = pil_img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                elif max_height:
                    ratio = max_height / pil_img.height
                    new_width = int(pil_img.width * ratio)
                    pil_img = pil_img.resize((new_width, max_height), Image.Resampling.LANCZOS)
            
            # Save processed image temporarily if needed
            if max_width or max_height or quality != 95:
                temp_file = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
                pil_img.save(temp_file.name, 'JPEG', quality=quality)
                processed_images.append(temp_file.name)
                temp_files.append(temp_file.name)
            else:
                processed_images.append(str(img_path_obj))
                
        except Exception as e:
            print(f"Error processing image {img_path}: {e}")
            continue
    
    # Calculate rows needed
    num_rows = (len(processed_images) + cols - 1) // cols  # Round up
    
    # Create table
    table = doc.add_table(rows=num_rows, cols=cols)
    table.style = 'Table Grid'
    
    # Set column widths
    for col in range(cols):
        table.columns[col].width = Inches(column_width / 1440 * 20)  # Convert to twips
    
    # Fill table with images and filenames
    for idx, img_path in enumerate(processed_images):
        row_idx = idx // cols
        col_idx = idx % cols
        
        if row_idx >= num_rows:
            break
        
        row = table.rows[row_idx]
        cell = row.cells[col_idx]
        cell.vertical_alignment = 1  # Center vertically
        
        try:
            # Get image dimensions maintaining aspect ratio
            img = Image.open(img_path)
            original_width, original_height = img.size
            
            aspect_ratio = original_height / original_width
            image_height = image_width * aspect_ratio
            
            # Limit height
            max_height = Inches(4.5)
            if image_height > max_height:
                image_height = max_height
                image_width_adjusted = image_height / aspect_ratio
            else:
                image_width_adjusted = image_width
            
            # Add image
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(img_path), width=image_width_adjusted, height=image_height)
            
            # Add filename below image
            filename = image_names[idx] if idx < len(image_names) else Path(img_path).stem
            filename_para = cell.add_paragraph()
            filename_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            filename_run = filename_para.add_run(filename)
            filename_run.font.size = Pt(10)
            filename_run.font.name = 'Arial'
            
        except Exception as e:
            print(f"Error adding image {img_path}: {e}")
            filename = image_names[idx] if idx < len(image_names) else Path(img_path).stem
            cell.text = f"Error loading: {filename}"
        
        # Set 1px black borders
        border_style = {"sz": "8", "val": "single", "color": "#000000"}
        set_cell_border(cell, top=border_style, left=border_style, bottom=border_style, right=border_style)
    
    # Save document
    doc.save(str(output_path))
    
    # Clean up temp files
    for temp_file in temp_files:
        try:
            os.unlink(temp_file)
        except:
            pass
    
    return str(output_path)

