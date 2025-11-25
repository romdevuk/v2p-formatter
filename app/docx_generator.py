from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
from pathlib import Path

# Convert 5px to twips (1 inch = 1440 twips, 1px ≈ 0.75pt, 1pt = 20 twips)
# 5px ≈ 3.75pt ≈ 75 twips
BOTTOM_PADDING_TWIPS = 75  # 5px spacing

def set_cell_border(cell, **kwargs):
    """
    Set cell border for a table cell
    Usage: set_cell_border(cell, top={"sz": 8, "val": "single", "color": "#000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_val = kwargs.get(border_name)
        if border_val:
            tag = 'w:{}'.format(border_name)
            element = OxmlElement(tag)
            for key, value in border_val.items():
                element.set(qn('w:{}'.format(key)), str(value))
            tcPr.append(element)

def create_docx(images, output_path, images_per_page=1):
    """
    Create DOCX document with images in a 2-column table layout
    
    Args:
        images: List of image file paths
        output_path: Path to save DOCX file
        images_per_page: Number of images per page (ignored, always 2 columns)
    """
    doc = Document()
    
    # Set default margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Calculate image size for 2 columns (with margins and spacing)
    # Page width: 8.27", margins: 2*0.5" = 1", usable: 7.27"
    # Two columns with spacing: each column ~3.5" wide
    column_width = Inches(3.5)
    image_width = Inches(3.3)  # Slightly smaller to fit in cell with padding
    
    # Calculate number of rows needed (2 images per row)
    num_rows = (len(images) + 1) // 2  # Round up
    
    # Create a single table with all rows
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = column_width
    table.columns[1].width = column_width
    
    # Fill table with images
    for i in range(0, len(images), 2):
        row_idx = i // 2
        row = table.rows[row_idx]
        
        # First column
        cell1 = row.cells[0]
        cell1.vertical_alignment = 1  # Center vertically
        if i < len(images):
            try:
                img = Image.open(images[i])
                original_width, original_height = img.size
                
                # Calculate height maintaining aspect ratio
                aspect_ratio = original_height / original_width
                image_height = image_width * aspect_ratio
                
                # Limit height to reasonable size
                max_height = Inches(4.5)
                if image_height > max_height:
                    image_height = max_height
                    image_width = image_height / aspect_ratio
                
                paragraph = cell1.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(images[i]), width=image_width, height=image_height)
            except Exception as e:
                print(f"Error adding image {images[i]}: {e}")
                cell1.text = f"Error loading image: {Path(images[i]).name}"
        
        # Second column
        cell2 = row.cells[1]
        cell2.vertical_alignment = 1  # Center vertically
        if i + 1 < len(images):
            try:
                img = Image.open(images[i + 1])
                original_width, original_height = img.size
                
                # Calculate height maintaining aspect ratio
                aspect_ratio = original_height / original_width
                image_height = image_width * aspect_ratio
                
                # Limit height to reasonable size
                max_height = Inches(4.5)
                if image_height > max_height:
                    image_height = max_height
                    image_width = image_height / aspect_ratio
                
                paragraph = cell2.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(images[i + 1]), width=image_width, height=image_height)
            except Exception as e:
                print(f"Error adding image {images[i + 1]}: {e}")
                cell2.text = f"Error loading image: {Path(images[i + 1]).name}"
        else:
            # Empty cell for odd number of images
            cell2.text = ""
        
        # Set 1px black borders on all cells
        border_style = {"sz": "8", "val": "single", "color": "#000000"}  # 8 = 1pt = 1px equivalent
        set_cell_border(cell1, top=border_style, left=border_style, bottom=border_style, right=border_style)
        set_cell_border(cell2, top=border_style, left=border_style, bottom=border_style, right=border_style)
        
        # Add 5px bottom padding to cells
        # Set cell padding using tcPr (table cell properties)
        for cell in [cell1, cell2]:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Create tcMar (table cell margins) element
            tcMar = OxmlElement('w:tcMar')
            
            # Bottom margin
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:w'), str(BOTTOM_PADDING_TWIPS))
            bottom.set(qn('w:type'), 'dxa')  # dxa = twentieths of a point (twips)
            tcMar.append(bottom)
            
            tcPr.append(tcMar)
    
    # Save document
    doc.save(str(output_path))
    return str(output_path)
