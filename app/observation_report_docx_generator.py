"""
Observation Report - DOCX Generator Module

Purpose: Generate DOCX files from draft data

⚠️ IMPORTANT: This is a NEW module - no legacy code from old modules.
The old observation-media module did not work properly and must be completely avoided.

Author: Backend Developer (Agent-1)
Created: Stage 1
"""

from pathlib import Path
from typing import Dict, List, Optional
import logging
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from PIL import Image
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.error("python-docx or PIL not available - DOCX generation will not work")

logger = logging.getLogger(__name__)

# Placeholder pattern
PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Za-z0-9_]+)\}\}')


def generate_docx(
    text_content: str,
    assignments: Dict,
    header_data: Dict,
    assessor_feedback: str,
    filename: str,
    output_folder: Path,
    font_size: int = 16,
    font_name: str = "Times New Roman"
) -> Path:
    """
    Generate DOCX file from draft data
    
    Args:
        text_content: Text content with placeholders
        assignments: Media assignments dictionary
        header_data: Header information dictionary
        assessor_feedback: Assessor feedback text
        filename: Output filename (without .docx extension)
        output_folder: Base output folder path
        font_size: Font size in points (default: 16)
        font_name: Font name (default: "Times New Roman")
        
    Returns:
        Path to generated DOCX file
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available")
    
    # Ensure filename has .docx extension
    if not filename.endswith('.docx'):
        filename = f"{filename}.docx"
    
    # Create new document
    doc = Document()
    
    # Set A4 page size
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)  # A4 height
        section.page_width = Inches(8.27)    # A4 width
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Add "Assessment Report" heading if header_data exists
    if header_data and any(header_data.values()):
        heading = doc.add_heading('Assessment Report', level=1)
        heading_format = heading.paragraph_format
        heading_format.space_after = Pt(12)
    
    # Add header table if header_data exists
    if header_data and any(header_data.values()):
        _add_header_table(doc, header_data, font_name, font_size)
        doc.add_paragraph()  # Blank paragraph
    
    # Add "Observation Report" heading if header_data exists
    if header_data and any(header_data.values()):
        heading = doc.add_heading('Observation Report', level=1)
        heading_format = heading.paragraph_format
        heading_format.space_after = Pt(12)
    
    # Process text content and replace placeholders with tables
    _process_text_content(doc, text_content, assignments, font_name, font_size)
    
    # Add assessor feedback table if feedback exists
    if assessor_feedback and assessor_feedback.strip():
        doc.add_paragraph()  # Blank paragraph
        _add_assessor_feedback_table(doc, assessor_feedback, font_name, font_size)
    
    # Save document
    output_path = output_folder / filename
    doc.save(str(output_path))
    logger.info(f"DOCX generated: {output_path}")
    
    return output_path


def _add_header_table(doc: Document, header_data: Dict, font_name: str, font_size: int):
    """Add header information as a 2-column table"""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.77)
    
    # Helper function to add border to cell
    def set_cell_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for border_name in ['top', 'left', 'bottom', 'right']:
            tag = f'w:{border_name}'
            element = OxmlElement(tag)
            element.set(qn('w:sz'), '4')  # 0.25pt
            element.set(qn('w:val'), 'single')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)
    
    # Add rows for each header field
    fields = ['learner', 'assessor', 'visit_date', 'location', 'address']
    labels = ['Learner', 'Assessor', 'Visit Date', 'Location', 'Address']
    
    for label, field in zip(labels, fields):
        value = header_data.get(field, '')
        if value:
            # Format visit_date if present
            if field == 'visit_date' and value:
                value = _format_date(value)
            
            row = table.add_row()
            # Label cell
            cell1 = row.cells[0]
            cell1.text = label
            set_cell_border(cell1)
            _set_cell_formatting(cell1, font_name, font_size, bold=True)
            
            # Value cell
            cell2 = row.cells[1]
            cell2.text = str(value)
            set_cell_border(cell2)
            _set_cell_formatting(cell2, font_name, font_size)


def _process_text_content(doc: Document, text_content: str, assignments: Dict, font_name: str, font_size: int):
    """Process text content and replace placeholders with 2-column media tables"""
    if not text_content:
        return
    
    # Split content by placeholders
    parts = PLACEHOLDER_PATTERN.split(text_content)
    
    current_text = ""
    
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Regular text (even indices)
            current_text += part
        else:
            # Placeholder name (odd indices) - output accumulated text, then placeholder table
            if current_text.strip():
                # Add accumulated text as paragraphs
                _add_text_as_paragraphs(doc, current_text.strip(), font_name, font_size)
                current_text = ""
            
            # Add placeholder table
            placeholder_name = part.lower()
            if placeholder_name in assignments:
                media_list = assignments[placeholder_name]
                _add_media_table(doc, media_list, font_name, font_size)
    
    # Add any remaining text
    if current_text.strip():
        _add_text_as_paragraphs(doc, current_text.strip(), font_name, font_size)


def _add_text_as_paragraphs(doc: Document, text: str, font_name: str, font_size: int):
    """Add text as paragraphs, handling sections"""
    # Split by lines
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            # Empty line - add blank paragraph
            doc.add_paragraph()
        elif line.upper().startswith('SECTION'):
            # Section heading
            section_title = line.replace('SECTION', '').replace(':', '').replace('-', '').strip()
            if section_title:
                heading = doc.add_heading(section_title, level=2)
                heading_format = heading.paragraph_format
                heading_format.space_after = Pt(6)
        else:
            # Regular paragraph
            para = doc.add_paragraph(line)
            _set_paragraph_formatting(para, font_name, font_size)


def _add_media_table(doc: Document, media_list: List[Dict], font_name: str, font_size: int):
    """Add 2-column table for media items"""
    if not media_list:
        return
    
    # Sort by order
    sorted_media = sorted(media_list, key=lambda x: x.get('order', 0))
    
    # Calculate number of rows (2 columns)
    num_rows = (len(sorted_media) + 1) // 2
    
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    
    # Set column widths
    column_width = Inches(3.5)
    table.columns[0].width = column_width
    table.columns[1].width = column_width
    
    # Set cell borders
    def set_cell_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for border_name in ['top', 'left', 'bottom', 'right']:
            tag = f'w:{border_name}'
            element = OxmlElement(tag)
            element.set(qn('w:sz'), '4')  # 0.25pt
            element.set(qn('w:val'), 'single')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)
    
    # Fill table
    for idx, media_item in enumerate(sorted_media):
        row_idx = idx // 2
        col_idx = idx % 2
        
        row = table.rows[row_idx]
        cell = row.cells[col_idx]
        set_cell_border(cell)
        
        media_path = Path(media_item.get('path', ''))
        media_type = media_item.get('type', '')
        
        if media_type == 'image' and media_path.exists():
            # Embed image
            try:
                img = Image.open(media_path)
                original_width, original_height = img.size
                aspect_ratio = original_height / original_width
                
                image_width = Inches(3.0)
                image_height = image_width * aspect_ratio
                max_height = Inches(4.0)
                if image_height > max_height:
                    image_height = max_height
                    image_width = image_height / aspect_ratio
                
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(media_path), width=image_width, height=image_height)
            except Exception as e:
                logger.warning(f"Error adding image {media_path}: {e}")
                cell.text = media_path.name
        else:
            # Add filename as text for videos, PDFs, MP3s
            cell.text = media_path.name if media_path.name else str(media_path)
            _set_cell_formatting(cell, font_name, font_size)
            cell.vertical_alignment = 1  # Center vertically


def _add_assessor_feedback_table(doc: Document, feedback: str, font_name: str, font_size: int):
    """Add assessor feedback as a table"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # Set cell width
    table.columns[0].width = Inches(7.27)
    
    # Set border
    cell = table.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'left', 'bottom', 'right']:
        tag = f'w:{border_name}'
        element = OxmlElement(tag)
        element.set(qn('w:sz'), '4')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:color'), '000000')
        tcPr.append(element)
    
    # Add feedback text
    cell.text = feedback
    _set_cell_formatting(cell, font_name, font_size)


def _set_cell_formatting(cell, font_name: str, font_size: int, bold: bool = False):
    """Set formatting for table cell"""
    for paragraph in cell.paragraphs:
        _set_paragraph_formatting(paragraph, font_name, font_size, bold)


def _set_paragraph_formatting(paragraph, font_name: str, font_size: int, bold: bool = False):
    """Set formatting for paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold


def _format_date(date_str: str) -> str:
    """Format date from YYYY-MM-DD to readable format"""
    try:
        from datetime import datetime
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        return dt.strftime('%d %B %Y')
    except (ValueError, TypeError):
        return date_str  # Return as-is if parsing fails

