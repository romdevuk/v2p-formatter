"""
Observation Media DOCX Generator
Generates DOCX files with embedded media in 2-column tables
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
from PIL import Image

logger = logging.getLogger(__name__)


def create_observation_docx(text_content: str, assignments: dict, output_path: Path, font_size: int = 16, font_name: str = 'Times New Roman', header_data: dict = None, assessor_feedback: str = None) -> dict:
    """
    Create a DOCX file from observation media text and assignments.
    
    Args:
        text_content: Text content with placeholders
        assignments: Dictionary mapping placeholder names to media lists
        output_path: Path where DOCX file should be saved
        font_size: Font size in points (default: 16)
        font_name: Font name (default: 'Times New Roman')
        header_data: Dictionary with header fields (learner, assessor, visit_date, location, address) (optional)
        assessor_feedback: Assessor feedback text (optional)
        
    Returns:
        Dictionary with success status and file path
    """
    logger.info(f"create_observation_docx called with {len(assignments)} placeholder assignments, font_size={font_size}, font_name={font_name}")
    try:
        # Create document
        doc = Document()
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        # Ensure document is editable (no protection)
        # Remove any document protection if it exists
        try:
            settings_part = doc.part.settings_part
            if settings_part is not None:
                settings = settings_part.settings
                # Remove document protection if present
                docProtection = settings.element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}documentProtection')
                if docProtection is not None:
                    settings.element.remove(docProtection)
        except:
            pass  # If settings can't be accessed, continue anyway
        
        # Set page size to A4
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4 height
        section.page_width = Inches(8.27)   # A4 width
        # Default margins are 1 inch on each side
        
        # Add header section if header_data is provided
        if header_data:
            # Format date from YYYY-MM-DD to readable format
            def format_date(date_string):
                if not date_string:
                    return ''
                try:
                    from datetime import datetime
                    date_obj = datetime.strptime(date_string, '%Y-%m-%d')
                    months = ['January', 'February', 'March', 'April', 'May', 'June',
                             'July', 'August', 'September', 'October', 'November', 'December']
                    return f"{date_obj.day} {months[date_obj.month - 1]} {date_obj.year}"
                except:
                    return date_string
            
            # Add "Assessment Report" heading
            heading1 = doc.add_paragraph('Assessment Report')
            heading1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading1_run = heading1.runs[0]
            heading1_run.font.name = font_name
            heading1_run.font.size = Pt(18)
            heading1_run.font.bold = True
            heading1.paragraph_format.space_after = Pt(12)
            
            # Add header table
            header_table = doc.add_table(rows=5, cols=2)
            header_table.style = 'Light Grid Accent 1'  # Basic table style
            
            # Set table borders to 1pt solid black (sz='8' = 1pt in Word XML)
            for row in header_table.rows:
                for cell in row.cells:
                    cell_borders = cell._element.get_or_add_tcPr().get_or_add_tcBorders()
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'{qn("w")}{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '8')  # 1pt (Word XML sz is in eighths of a point)
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), '000000')
                        cell_borders.append(border)
            
            # Fill header table
            header_rows = [
                ('Learner', header_data.get('learner', '')),
                ('Assessor', header_data.get('assessor', '')),
                ('Visit Date', format_date(header_data.get('visit_date', ''))),
                ('Location', header_data.get('location', '')),
                ('Address', header_data.get('address', ''))
            ]
            
            for idx, (label, value) in enumerate(header_rows):
                header_table.rows[idx].cells[0].text = label
                header_table.rows[idx].cells[1].text = value
                
                # Set font for both cells
                for cell in header_table.rows[idx].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
            
            # Set first column width
            header_table.columns[0].width = Inches(1.5)
            header_table.columns[1].width = Inches(5.5)
            
            # Add spacing after header table
            doc.add_paragraph()  # Blank paragraph
            
            # Add "Observation Report" heading
            heading2 = doc.add_paragraph('Observation Report')
            heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading2_run = heading2.runs[0]
            heading2_run.font.name = font_name
            heading2_run.font.size = Pt(18)
            heading2_run.font.bold = True
            heading2.paragraph_format.space_before = Pt(12)
            heading2.paragraph_format.space_after = Pt(12)
        
        # Parse text and replace placeholders with tables
        import re
        
        # If no text content, add a default paragraph
        if not text_content or not text_content.strip():
            doc.add_paragraph("No content")
        else:
            lines = text_content.split('\n')
            
            for line_idx, line in enumerate(lines):
                # Handle empty lines
                if not line.strip():
                    # Empty paragraphs will inherit font from Normal style we set earlier
                    doc.add_paragraph()
                    continue
                
                # Check if line contains placeholders
                placeholder_pattern = re.compile(r'(\{\{[A-Za-z0-9_]+\}\})')
                placeholders_found = placeholder_pattern.findall(line)
                
                if placeholders_found:
                    # Split line by placeholders while keeping them
                    parts = placeholder_pattern.split(line)
                    
                    current_paragraph = None
                    for part in parts:
                        if not part:
                            continue
                        
                        # Check if this part is a placeholder
                        placeholder_match = placeholder_pattern.match(part)
                        if placeholder_match:
                            # Save any current paragraph before adding table
                            current_paragraph = None
                            
                            # Extract placeholder name
                            # The pattern matches {{Name}}, so group(0) is the full match including braces
                            # We need to extract just the name inside the braces
                            full_match = placeholder_match.group(0)  # e.g., '{{Image1}}'
                            # Remove braces: {{Image1}} -> Image1
                            placeholder_name = full_match.strip('{}')
                            placeholder_key = placeholder_name.lower()
                            assigned_media = assignments.get(placeholder_key, [])
                            
                            logger.info(f"Processing placeholder '{placeholder_key}': {len(assigned_media)} media items")
                            
                            # Add table for placeholder
                            _add_media_table_to_doc(doc, assigned_media)
                        else:
                            # Regular text - add to current paragraph or create new one
                            if not current_paragraph:
                                current_paragraph = doc.add_paragraph()
                            run = current_paragraph.add_run(part)
                            # Apply font settings to run
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
                else:
                    # No placeholders - just add the text as a paragraph
                    if line.strip():
                        para = doc.add_paragraph(line.strip())
                        # Apply font settings to all runs in paragraph
                        for run in para.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
        
        # Ensure document has at least one element (paragraph or table)
        # Check both paragraphs and if document has any tables
        has_content = len(doc.paragraphs) > 0 or len(doc.tables) > 0
        if not has_content:
            doc.add_paragraph("Empty document")
        
        # Verify we have content before saving
        has_paragraphs = len(doc.paragraphs) > 0
        has_tables = len(doc.tables) > 0
        
        if not has_paragraphs and not has_tables:
            logger.warning("Document has no content, adding default paragraph")
            doc.add_paragraph("Document content")
        
        # Add assessor feedback table at the bottom if provided
        if assessor_feedback:
            # Add blank paragraph before assessor feedback
            doc.add_paragraph()
            
            # Add assessor feedback table
            feedback_table = doc.add_table(rows=2, cols=1)
            feedback_table.style = 'Light Grid Accent 1'
            
            # Set table borders to 1pt solid black
            for row in feedback_table.rows:
                for cell in row.cells:
                    cell_borders = cell._element.get_or_add_tcPr().get_or_add_tcBorders()
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'{qn("w")}{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '8')  # 1pt
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), '000000')
                        cell_borders.append(border)
            
            # Fill assessor feedback table
            feedback_table.rows[0].cells[0].text = 'Assessor Feedback'
            feedback_table.rows[1].cells[0].text = assessor_feedback
            
            # Set font for both cells
            for row in feedback_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run.font.size = Pt(font_size)
            
            # Set table width
            feedback_table.columns[0].width = Inches(7)
        
        # Ensure output path has .docx extension
        if not str(output_path).lower().endswith('.docx'):
            output_path = output_path.with_suffix('.docx')
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        # Verify file was created and has size > 0
        if output_path.exists():
            file_size = output_path.stat().st_size
            logger.info(f"DOCX file created successfully: {output_path} ({file_size} bytes)")
            
            if file_size == 0:
                logger.error("DOCX file is empty after save!")
                return {
                    'success': False,
                    'error': 'Generated DOCX file is empty'
                }
        else:
            logger.error(f"DOCX file was not created: {output_path}")
            return {
                'success': False,
                'error': 'Failed to create DOCX file'
            }
        
        return {
            'success': True,
            'file_path': str(output_path),
            'file_name': output_path.name
        }
        
    except Exception as e:
        logger.error(f"Error creating DOCX file: {e}", exc_info=True)
        return {
            'success': False,
            'error': str(e)
        }


def _extract_placeholders_from_line(line: str) -> list:
    """Extract placeholder names from a line of text"""
    import re
    pattern = re.compile(r'\{\{([A-Za-z0-9_]+)\}\}')
    matches = pattern.findall(line)
    return [m.lower() for m in matches]


def _add_media_table_to_doc(doc: Document, media_list: list):
    """
    Add a 2-column table with media to the document.
    Simple approach: equal column widths (50% each), just like the preview.
    
    Args:
        doc: Document object
        media_list: List of media dictionaries with 'path', 'name', 'type'
    """
    # Get page width and margins to calculate available width
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    
    # Convert Length objects to inches for calculation
    try:
        page_width_inches = float(page_width)
        left_margin_inches = float(left_margin)
        right_margin_inches = float(right_margin)
        available_width_inches = page_width_inches - left_margin_inches - right_margin_inches
    except (ValueError, TypeError):
        # Fallback: A4 page width (8.27") minus margins (1" each side) = 6.27"
        available_width_inches = 6.27
    
    # Calculate column width - half of available width
    column_width_inches = available_width_inches / 2.0
    column_width = Inches(column_width_inches)
    
    # Create available_width Length object for table width setting
    available_width = Inches(available_width_inches)
    
    if not media_list:
        # Empty table
        logger.debug("Creating empty table (no media)")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        _style_table(table)
        _set_table_width(table, available_width)
        row = table.rows[0]
        _style_cell(row.cells[0])
        _style_cell(row.cells[1])
        # Set equal column widths
        table.columns[0].width = column_width
        table.columns[1].width = column_width
        return
    
    logger.info(f"_add_media_table_to_doc: Adding {len(media_list)} media items")
    
    # Calculate number of rows needed (2 columns per row)
    num_rows = (len(media_list) + 1) // 2
    
    # Create table
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    _style_table(table)
    _set_table_width(table, available_width)
    
    # Set equal column widths (50% each) - simple like the preview
    # Set widths on columns - this is the primary method
    table.columns[0].width = column_width
    table.columns[1].width = column_width
    
    # Style all cells (borders, etc.) but don't override cell widths
    # Let the column widths control the layout
    for row in table.rows:
        for cell in row.cells:
            _style_cell(cell)
    
    # Fill table with media
    logger.info(f"Adding {len(media_list)} media items to table")
    for i, media in enumerate(media_list):
        row_idx = i // 2
        col_idx = i % 2
        
        cell = table.rows[row_idx].cells[col_idx]
        
        # Add image or video to cell
        if media['type'] == 'image':
            logger.info(f"Processing image {i+1}/{len(media_list)}: {media.get('name', 'unknown')} at path: {media.get('path', 'unknown')}")
            # Add image
            try:
                # Resolve path - paths from scanner are absolute, but might be URLs when sent from frontend
                media_path = media['path']
                logger.debug(f"Resolving image path: {media_path}")
                image_path = None
                
                # First, try as absolute path (most common case from observation_media_scanner)
                test_path = Path(media_path)
                if test_path.is_absolute() and test_path.exists():
                    image_path = test_path
                elif test_path.is_absolute():
                    # Absolute path but doesn't exist - might be on different system or moved
                    logger.warning(f"Absolute path doesn't exist: {media_path}")
                else:
                    # Not absolute - might be a relative path or URL
                    # Check if it's a static URL
                    if '/static/' in media_path or media_path.startswith('/static/'):
                        # Extract relative path from static URL
                        parts = media_path.split('/static/', 1)
                        if len(parts) > 1:
                            rel_path = parts[1]
                            # Try Flask static folder
                            from flask import current_app
                            try:
                                static_folder = current_app.static_folder
                                image_path = Path(static_folder) / rel_path
                            except:
                                # Fallback
                                base_dir = Path(__file__).parent.parent
                                image_path = base_dir / 'static' / rel_path
                    else:
                        # Try as relative path
                        base_dir = Path(__file__).parent.parent
                        image_path = base_dir / media_path.lstrip('/')
                
                # If still not found, check OUTPUT_FOLDER (where observation media scanner looks)
                if not image_path or not image_path.exists():
                    from config import OUTPUT_FOLDER
                    # Try direct filename match in OUTPUT_FOLDER
                    filename = Path(media_path).name
                    potential_path = OUTPUT_FOLDER / filename
                    if potential_path.exists():
                        image_path = potential_path
                        logger.info(f"Found image in OUTPUT_FOLDER: {image_path}")
                    else:
                        # Try recursive search in OUTPUT_FOLDER
                        matches = list(OUTPUT_FOLDER.rglob(filename))
                        if matches:
                            image_path = matches[0]
                            logger.info(f"Found image by name search: {image_path}")
                
                # Log if still not found
                if not image_path or not image_path.exists():
                    logger.error(f"Could not resolve image path: {media_path}")
                    logger.error(f"  Tried: {image_path if image_path else 'None'}")
                    # Try one more time - check if it's just the filename and search OUTPUT_FOLDER recursively
                    if image_path:
                        filename_only = Path(media_path).name
                        from config import OUTPUT_FOLDER
                        all_matches = list(OUTPUT_FOLDER.rglob(filename_only))
                        if all_matches:
                            image_path = all_matches[0]
                            logger.info(f"Found image by recursive filename search: {image_path}")
                
                if image_path and image_path.exists():
                    # Get image dimensions
                    img = Image.open(image_path)
                    width, height = img.size
                    
                    # Calculate image size to fit in cell (similar to preview)
                    # Use column width minus padding for max width
                    max_width_inches = column_width_inches - 0.2  # 0.1 inch padding each side
                    max_width = Inches(max_width_inches)
                    max_height = Inches(4.5)  # Reasonable max height
                    
                    # Calculate aspect ratio
                    aspect_ratio = height / width
                    
                    # Fit to width first
                    image_width = max_width
                    image_height = image_width * aspect_ratio
                    
                    # If height exceeds max, scale down
                    if image_height > max_height:
                        image_height = max_height
                        image_width = image_height / aspect_ratio
                    
                    # Add image to cell
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    try:
                        run.add_picture(str(image_path), width=image_width, height=image_height)
                        logger.info(f"✓ Added image to DOCX: {image_path}")
                    except Exception as pic_error:
                        logger.error(f"Error adding picture to DOCX: {pic_error}", exc_info=True)
                        raise
                else:
                    # Image not found - add text placeholder
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(f"[Image not found: {media['name']}]")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    logger.warning(f"Image not found: original path={media_path}, resolved={image_path}")
            except Exception as e:
                logger.warning(f"Error adding image {media['path']}: {e}", exc_info=True)
                paragraph = cell.paragraphs[0]
                paragraph.add_run(f"[Error loading image: {media['name']}]")
        elif media['type'] == 'document':
            # PDF - just add filename (like video)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(media['name'])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # Video - just add filename
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(media['name'])
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Style empty cells if there's an odd number of images
    # (The last row will have one empty cell)
    if len(media_list) % 2 == 1:
        last_row = table.rows[-1]
        empty_cell = last_row.cells[1]
        _style_cell(empty_cell)
        # Empty cell - leave it empty


def _style_table(table):
    """Apply styling to table (black borders, 1px)"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tbl.tblPr.append(tblBorders)


def _set_table_width(table, width):
    """Set table width to specified value (width should be a docx.shared.Length object)"""
    from docx.shared import Length
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        # Remove existing width if present
        existing_width = tblPr.find(qn('w:tblW'))
        if existing_width is not None:
            tblPr.remove(existing_width)
    
    # Convert width to twips
    # 1 inch = 914400 EMU = 1440 twips
    # Best approach: convert Length to inches first, then to twips
    try:
        if isinstance(width, Length):
            # Try to get inches directly from Length object
            try:
                # Length objects in python-docx have an __float__ method that returns inches
                width_inches = float(width)
            except (ValueError, TypeError):
                # Fallback: convert from EMU
                width_emu = width.emu
                width_inches = width_emu / 914400.0
            
            # Convert inches to twips: 1 inch = 1440 twips
            width_twips = int(round(width_inches * 1440))
        else:
            # If it's a number, assume it's already in inches
            width_twips = int(round(float(width) * 1440))
    except (ValueError, TypeError, AttributeError) as e:
        # Fallback: calculate from available_width_inches if we have it
        # For A4 with 1 inch margins: 8.27 - 2 = 6.27 inches = 9029 twips
        logger.warning(f"Error converting width to twips: {e}, using default")
        width_twips = 9000  # Default: ~6.25 inches (reasonable for A4 with margins)
    
    # Sanity check - table width should be reasonable (between 5000-10000 twips for A4)
    if width_twips < 1000 or width_twips > 20000:
        logger.warning(f"Table width seems incorrect ({width_twips} twips), using default")
        width_twips = 9000  # Default: ~6.25 inches
    
    # Set table width
    tblWidth = OxmlElement('w:tblW')
    tblWidth.set(qn('w:w'), str(width_twips))
    tblWidth.set(qn('w:type'), 'dxa')  # Width in twips (1/20th of a point)
    tblPr.append(tblWidth)
    
    # Set table to use fixed layout to respect width
    existing_layout = tblPr.find(qn('w:tblLayout'))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)


def _set_table_alignment(table, alignment='left'):
    """Set table alignment (left, center, right)"""
    # alignment can be 'left', 'center', 'right'
    alignment_map = {
        'left': 'left',
        'center': 'center',
        'right': 'right'
    }
    
    if alignment not in alignment_map:
        alignment = 'left'
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing alignment if present
    existing_jc = tblPr.find(qn('w:jc'))
    if existing_jc is not None:
        tblPr.remove(existing_jc)
    
    # Set table justification
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), alignment_map[alignment])
    tblPr.append(jc)


def _style_cell(cell):
    """Apply styling to cell (black borders, padding)"""
    cell.vertical_alignment = 1  # Center vertically
    # Borders are handled by table styling
    # Add some padding
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._element.append(tcPr)
    
    tcMar = OxmlElement('w:tcMar')
    for margin_name in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), '144')  # 0.1 inch
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    
    tcPr.append(tcMar)

